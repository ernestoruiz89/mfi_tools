import io
import re
from html import unescape
from html.parser import HTMLParser

import frappe
from frappe import _
from frappe.utils import cint, cstr, flt, now_datetime
from mfi_tools.mfi_tools.utils.estado_line_format import format_accounting_number, format_estado_line_value, is_text_estado_line
from mfi_tools.mfi_tools.utils.nota_eeff import get_package_note_rows
from mfi_tools.mfi_tools.utils.nota_tablas import build_complex_section_tables
from mfi_tools.mfi_tools.utils.customer import get_customer_display

REPORT_TITLE = "Estados Financieros y Notas"
FONT_NAME = "Arial Narrow"
BODY_SIZE = 12
COMPACT_SIZE = 12
COMPLEX_NOTE_TABLE_SIZE = 10


def export_paquete_eeff_to_word(package_name):
    file_name, content = build_paquete_eeff_word_content(package_name)

    from frappe.utils.file_manager import save_file

    file_doc = save_file(file_name, content, "Paquete EEFF", package_name, is_private=1)
    return {
        "file_name": file_doc.file_name,
        "file_url": file_doc.file_url,
        "file_id": file_doc.name,
    }


def build_paquete_eeff_word_content(package_name):
    package = frappe.get_doc("Paquete EEFF", package_name)
    document = _build_package_document(package)

    stream = io.BytesIO()
    document.save(stream)
    content = stream.getvalue()
    stream.close()

    file_name = _build_word_export_filename(package.name)
    return file_name, content


def _build_package_document(package):
    Document, WD_ALIGN_PARAGRAPH, WD_ORIENTATION, WD_SECTION_START, WD_TABLE_ALIGNMENT, OxmlElement, qn, Cm, Pt, RGBColor = _docx_imports()

    customer_display = get_customer_display(package.cliente)

    document = Document()
    _configure_document(document, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, Pt, document_title=REPORT_TITLE)

    title = document.add_paragraph(REPORT_TITLE, style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(customer_display or package.cliente or "Cliente")
    _set_paragraph_runs_font(subtitle)

    badge = document.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge.add_run("Documento preparado para emision de estados financieros").italic = True
    _set_paragraph_runs_font(badge)

    cover_table = document.add_table(rows=4, cols=2)
    cover_pairs = [
        ("Paquete", package.name or "-"),
        ("Cliente", customer_display or package.cliente or "-"),
        ("Periodo", package.periodo_nombre or "-"),
        ("Mes", package.mes or "-"),
        ("Anio", cstr(package.anio or "-")),
        ("Estado Preparacion", package.estado_preparacion or "-"),
        ("Balanza", package.balanza_comprobacion_eeff or "-"),
        ("Fecha Emision", cstr(package.fecha_emision or "-")),
    ]
    for index, pair in enumerate(cover_pairs):
        row = cover_table.rows[index // 2]
        cell = row.cells[index % 2]
        _fill_label_value_cell(cell, pair[0], pair[1])
    _style_meta_table(cover_table)

    summary = document.add_paragraph(
        "Se presenta el juego completo de estados financieros y notas explicativas para el periodo indicado."
    )
    _set_paragraph_runs_font(summary)

    toc_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    _configure_section(toc_section, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, landscape=False, document_title=REPORT_TITLE)
    _set_section_header_content(
        toc_section,
        {
            "cliente": customer_display or package.cliente or "Cliente",
            "titulo": "Estados Financieros",
            "periodo": "",
            "subtitulo": "",
        },
    )
    toc_heading = document.add_paragraph("Indice", style="Heading 1")
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_paragraph = document.add_paragraph()
    _append_field(toc_paragraph, 'TOC \\o "1-3" \\h \\z \\u', "Actualice el indice al abrir el documento.", OxmlElement, qn)
    content_section = document.add_section(WD_SECTION_START.NEW_PAGE)
    _configure_section(content_section, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, landscape=False, document_title=REPORT_TITLE)
    _set_section_footer_page_number(content_section, start=1)
    _add_estados_section(document, package)
    document.add_page_break()
    _add_notas_section(document, package)

    return document


def _docx_imports():
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except ImportError:
        frappe.throw(
            _(
                "No esta instalada la dependencia opcional python-docx. "
                "Instala la dependencia con <b>bench pip install python-docx</b> para habilitar la exportacion Word."
            ),
            title=_("Dependencia Faltante"),
        )
    return Document, WD_ALIGN_PARAGRAPH, WD_ORIENTATION, WD_SECTION_START, WD_TABLE_ALIGNMENT, OxmlElement, qn, Cm, Pt, RGBColor


def _configure_document(document, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, Pt, document_title=REPORT_TITLE):
    styles = document.styles
    styles["Normal"].font.name = FONT_NAME
    styles["Normal"].font.size = Pt(BODY_SIZE)
    styles["Title"].font.name = FONT_NAME
    styles["Title"].font.size = Pt(BODY_SIZE)
    styles["Heading 1"].font.name = FONT_NAME
    styles["Heading 1"].font.size = Pt(BODY_SIZE)
    styles["Heading 2"].font.name = FONT_NAME
    styles["Heading 2"].font.size = Pt(BODY_SIZE)
    styles["Heading 3"].font.name = FONT_NAME
    styles["Heading 3"].font.size = Pt(BODY_SIZE)
    if "Subtitle" in styles:
        styles["Subtitle"].font.name = FONT_NAME
        styles["Subtitle"].font.size = Pt(BODY_SIZE)

    for section in document.sections:
        _configure_section(section, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, landscape=False, document_title=document_title)


def _configure_section(section, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, landscape=False, document_title=REPORT_TITLE):
    _Document, _Align, WD_ORIENTATION, _SectionStart, _TableAlign, _OxmlElement, _qn, _Cm, _Pt, _RGBColor = _docx_imports()
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.2)
    section.orientation = WD_ORIENTATION.LANDSCAPE if landscape else WD_ORIENTATION.PORTRAIT
    if landscape:
        section.page_width, section.page_height = section.page_height, section.page_width

    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.text = ""
    _set_paragraph_runs_font(header_para, size=BODY_SIZE)

    footer = section.footer
    _clear_header_footer(footer)
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = ""
    _set_paragraph_runs_font(footer_para, size=BODY_SIZE)


def _append_field(paragraph, instruction, placeholder, OxmlElement, qn):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def _set_section_footer_page_number(section, start=None):
    Document, WD_ALIGN_PARAGRAPH, WD_ORIENTATION, WD_SECTION_START, WD_TABLE_ALIGNMENT, OxmlElement, qn, Cm, Pt, RGBColor = _docx_imports()
    footer = section.footer
    footer.is_linked_to_previous = False
    _clear_header_footer(footer)
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = ""
    _append_field(paragraph, "PAGE", "1", OxmlElement, qn)
    _set_paragraph_runs_font(paragraph, size=BODY_SIZE)

    if start is not None:
        sect_pr = section._sectPr
        pg_num = sect_pr.find(qn("w:pgNumType"))
        if pg_num is None:
            pg_num = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num)
        pg_num.set(qn("w:start"), cstr(start))


def _clear_header_footer(container):
    for paragraph in list(container.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)
    if not container.paragraphs:
        container.add_paragraph()


def _set_section_header_content(section, header_data):
    Document, WD_ALIGN_PARAGRAPH, WD_ORIENTATION, WD_SECTION_START, WD_TABLE_ALIGNMENT, OxmlElement, qn, Cm, Pt, RGBColor = _docx_imports()
    header = section.header
    _clear_header_footer(header)

    lines = [
        (cstr(header_data.get("cliente") or "-"), True),
        (cstr(header_data.get("titulo") or "-"), True),
    ]
    if cstr(header_data.get("periodo") or "").strip():
        lines.append((cstr(header_data.get("periodo") or ""), True))
    if cstr(header_data.get("subtitulo") or "").strip():
        lines.append((cstr(header_data.get("subtitulo") or ""), False))

    for index, (text, bold) in enumerate(lines):
        paragraph = header.paragraphs[index] if index < len(header.paragraphs) else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.text = cstr(text or "")
        _set_paragraph_runs_font(paragraph, size=BODY_SIZE, bold=bold)

    divider = header.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.LEFT
    divider.text = ""
    _set_paragraph_runs_font(divider, size=BODY_SIZE)
    for run in divider.runs:
        run.text = ""
    pPr = divider._element.get_or_add_pPr()
    pbdr = pPr.first_child_found_in("w:pBdr")
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pbdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "000000")


def _add_estados_section(document, package):
    document.add_paragraph("Estados Financieros", style="Heading 1")
    estados = frappe.get_all(
        "Estado Financiero EEFF",
        filters={"paquete_eeff": package.name},
        fields=["name", "tipo_estado", "titulo", "orden_presentacion"],
        order_by="orden_presentacion asc, creation asc",
        limit_page_length=100,
    )
    if not estados:
        paragraph = document.add_paragraph("No hay estados financieros registrados para este paquete.")
        _set_paragraph_runs_font(paragraph)
        return

    labels = package.get_column_labels()
    currency_symbol = _get_package_currency_symbol(package)

    Document, WD_ALIGN_PARAGRAPH, WD_ORIENTATION, WD_SECTION_START, WD_TABLE_ALIGNMENT, OxmlElement, qn, Cm, Pt, RGBColor = _docx_imports()

    for index, estado in enumerate(estados):
        estado_doc = frappe.get_doc("Estado Financiero EEFF", estado.name)
        is_landscape = cstr(estado_doc.get("orientacion") or "Vertical") == "Horizontal"
        
        if index == 0:
            section = document.sections[-1]
            if is_landscape:
                _configure_section(section, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, landscape=True, document_title=REPORT_TITLE)
        else:
            section = document.add_section(WD_SECTION_START.NEW_PAGE)
            _configure_section(section, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, landscape=is_landscape, document_title=REPORT_TITLE)
            _set_section_footer_page_number(section)
            
        header = estado_doc.get_print_header()
        _set_section_header_content(section, header)
        estado_font_size = estado_doc.get_print_font_size()

        if index == 0:
            document.add_paragraph("")

        table = document.add_table(rows=1, cols=5)
        _set_word_table_alignment(table, estado_doc.get_print_table_alignment())
        headers = ["", "Nota", labels["actual"], "", labels["comparativo"]]
        for header_index, title in enumerate(headers):
            table.rows[0].cells[header_index].text = title

        total_rows = []
        subtotal_rows = []
        rendered_rows = []

        lineas = list(estado_doc.lineas or [])
        for linea in lineas:
            if cint(getattr(linea, "no_imprimir", 0)):
                continue
            line_index = len(table.rows)
            row = table.add_row().cells
            row[0].text = "" if cint(getattr(linea, "es_linea_blanco", 0)) else (("    " * max((cint(linea.nivel or 1) - 1), 0)) + (cstr(linea.descripcion or "-")))
            if cint(getattr(linea, "es_linea_blanco", 0)):
                row[1].text = ""
                row[2].text = ""
                row[3].text = ""
                row[4].text = ""
            elif cint(getattr(linea, "es_titulo", 0)):
                row[1].text = cstr(getattr(linea, "nota", "") or "")
                row[2].text = ""
                row[3].text = ""
                row[4].text = ""
            elif is_text_estado_line(linea):
                row[1].text = cstr(getattr(linea, "nota", "") or "")
                row[2].text = format_estado_line_value(linea, "monto_actual")
                row[3].text = ""
                row[4].text = ""
            else:
                row[1].text = cstr(getattr(linea, "nota", "") or "")
                actual_value = format_estado_line_value(linea, "monto_actual")
                comparative_value = format_estado_line_value(linea, "monto_comparativo")
                if _is_estado_currency_line(linea):
                    actual_value = _apply_currency_symbol(actual_value, currency_symbol)
                    comparative_value = _apply_currency_symbol(comparative_value, currency_symbol)
                row[2].text = actual_value
                row[3].text = ""
                row[4].text = comparative_value
            rendered_rows.append((line_index, linea))
            if cint(getattr(linea, "es_total", 0)):
                total_rows.append(line_index)
            elif cint(getattr(linea, "es_subtotal", 0)):
                subtotal_rows.append(line_index)

        _set_table_column_widths(table, _get_estado_table_widths_cm(section, estado_doc))
        _style_financial_table(
            table,
            total_rows=total_rows,
            subtotal_rows=subtotal_rows,
            font_size=estado_font_size,
            note_col_index=1,
            numeric_col_indexes=(2, 4),
            gap_col_index=3,
            align_numeric_headers=True,
        )
        for line_index, linea in rendered_rows:
            _apply_estado_line_format(table.rows[line_index], linea, font_size=estado_font_size)
            if cint(getattr(linea, "es_linea_blanco", 0)):
                _set_row_min_height(table.rows[line_index], 340)

        if _package_has_signatures(package):
            _add_package_signatures_block(document, package)


def _apply_estado_line_format(row, linea, font_size=BODY_SIZE):
    desc_paragraph = row.cells[0].paragraphs[0]
    _set_paragraph_runs_font(
        desc_paragraph,
        bold=bool(
            cint(getattr(linea, "negrita", 0))
            or cint(getattr(linea, "es_total", 0))
            or cint(getattr(linea, "es_subtotal", 0))
        ),
    )
    if cint(getattr(linea, "subrayado", 0)):
        for run in desc_paragraph.runs:
            run.underline = True
    if cint(getattr(linea, "es_linea_blanco", 0)):
        return
    if is_text_estado_line(linea) and not cint(getattr(linea, "es_titulo", 0)):
        value_paragraph = row.cells[2].paragraphs[0]
        value_paragraph.alignment = 0
        _set_paragraph_runs_font(value_paragraph, size=font_size)


def _add_notas_section(document, package):
    document.add_paragraph("Notas a los Estados Financieros", style="Heading 1")
    labels = package.get_column_labels()
    currency_symbol = _get_package_currency_symbol(package)
    note_rows = get_package_note_rows(package.name, fields=["titulo", "sub_nota"], limit_page_length=300)
    if not note_rows:
        paragraph = document.add_paragraph("No hay notas registradas para este paquete.")
        _set_paragraph_runs_font(paragraph)
        return

    grouped_notes = {}
    ordered_numbers = []
    for note_row in note_rows:
        note_doc = frappe.get_doc("Nota EEFF", note_row.name)
        number_key = cint(getattr(note_doc, "numero_nota", 0) or 0)
        if number_key not in grouped_notes:
            grouped_notes[number_key] = {"principal": None, "subnotes": []}
            ordered_numbers.append(number_key)
        if cstr(getattr(note_doc, "sub_nota", "") or "").strip():
            grouped_notes[number_key]["subnotes"].append(note_doc)
        else:
            grouped_notes[number_key]["principal"] = note_doc

    printable_notes = []
    for number_key in ordered_numbers:
        group = grouped_notes[number_key]
        principal = group["principal"]
        subnotes = group["subnotes"]
        if principal:
            printable_notes.append((principal, subnotes))
            continue
        for subnote_doc in subnotes:
            printable_notes.append((subnote_doc, []))

    for note_index, (nota_doc, subnotes) in enumerate(printable_notes):
        if note_index > 0:
            document.add_page_break()
        _render_note_block(document, nota_doc, labels, package, currency_symbol, subnotes=subnotes)


def _render_note_block(document, nota_doc, labels, package, currency_symbol, subnotes=None, heading_style="Heading 2"):
    note_font_size = nota_doc.get_print_font_size() if hasattr(nota_doc, "get_print_font_size") else BODY_SIZE
    is_subnote = bool(cstr(getattr(nota_doc, "sub_nota", "") or "").strip())
    heading_paragraph = document.add_paragraph(
        cstr(nota_doc.get_print_heading() if hasattr(nota_doc, "get_print_heading") else cstr(nota_doc.titulo or "Sin titulo")),
        style=heading_style,
    )
    _set_paragraph_runs_font(heading_paragraph, size=BODY_SIZE)
    for run in heading_paragraph.runs:
        run.bold = not is_subnote

    if cstr(getattr(nota_doc, "estructura_nota", "Simple") or "Simple").strip() == "Compleja":
        _render_complex_note_content(document, nota_doc, labels, package, currency_symbol)
    else:
        _render_simple_note_content(document, nota_doc, labels, currency_symbol)

    rendered_observaciones = _get_rendered_note_observaciones(nota_doc)
    if rendered_observaciones:
        spacer = document.add_paragraph(" ")
        _set_paragraph_runs_font(spacer, size=BODY_SIZE)
        _add_rich_block(document, rendered_observaciones, size=BODY_SIZE)

    if subnotes:
        for subnote_doc in subnotes:
            spacer = document.add_paragraph(" ")
            _set_paragraph_runs_font(spacer, size=BODY_SIZE)
            _render_note_block(
                document,
                subnote_doc,
                labels,
                package,
                currency_symbol,
                subnotes=[],
                heading_style="Heading 3",
            )


def _render_simple_note_content(document, nota_doc, labels, currency_symbol):
    note_font_size = nota_doc.get_print_font_size() if hasattr(nota_doc, "get_print_font_size") else BODY_SIZE
    note_alignment = nota_doc.get_print_table_alignment() if hasattr(nota_doc, "get_print_table_alignment") else "Centro"
    rendered_narrative = _get_rendered_note_narrative(nota_doc)
    if rendered_narrative:
        _add_rich_block(document, rendered_narrative, size=BODY_SIZE)
        spacer = document.add_paragraph(" ")
        _set_paragraph_runs_font(spacer, size=BODY_SIZE)

    _render_note_figures(document, nota_doc, labels, note_font_size, note_alignment, currency_symbol)


def _render_note_figures(document, nota_doc, labels, note_font_size, note_alignment, currency_symbol):
    cifras = sorted(list(nota_doc.cifras_nota or []), key=lambda row: cint(row.idx or 0))
    visible_cifras = [row for row in cifras if not cint(getattr(row, "no_imprimir", 0))]
    if not visible_cifras:
        return

    table = document.add_table(rows=1, cols=4)
    _set_word_table_alignment(table, note_alignment)
    headers = ["", labels["actual"], "", labels["comparativo"]]
    for header_index, title in enumerate(headers):
        table.rows[0].cells[header_index].text = title

    total_actual = 0.0
    total_comparativo = 0.0
    total_rows = []
    subtotal_rows = []
    rendered_rows = []
    has_explicit_total = False
    total_format = _get_note_figures_total_format(nota_doc, visible_cifras)

    for cifra in visible_cifras:
        line_index = len(table.rows)
        row = table.add_row().cells
        concept_parts = []
        if not cint(getattr(cifra, "es_linea_blanco", 0)):
            indent = "    " * max(cint(getattr(cifra, "nivel", 1) or 1) - 1, 0)
            concept_parts.append(f"{indent}{cstr(cifra.concepto or '-')}")
        if cifra.comentario and not cint(getattr(cifra, "es_linea_blanco", 0)):
            concept_parts.append(cstr(cifra.comentario))
        row[0].text = "\n".join([part for part in concept_parts if part])
        if cint(getattr(cifra, "es_linea_blanco", 0)) or cint(getattr(cifra, "es_titulo", 0)):
            row[1].text = ""
            row[2].text = ""
            row[3].text = ""
        else:
            row[1].text = _format_note_figure_value(nota_doc, cifra, "monto_actual", currency_symbol)
            row[2].text = ""
            row[3].text = _format_note_figure_value(nota_doc, cifra, "monto_comparativo", currency_symbol)
        rendered_rows.append((line_index, cifra))

        if cint(getattr(cifra, "es_total", 0)):
            total_rows.append(line_index)
            has_explicit_total = True
        elif cint(getattr(cifra, "es_subtotal", 0)):
            subtotal_rows.append(line_index)
        elif not cint(getattr(cifra, "es_linea_blanco", 0)) and not cint(getattr(cifra, "es_titulo", 0)) and not _is_text_note_figure(nota_doc, cifra):
            total_actual += flt(cifra.monto_actual or 0)
            total_comparativo += flt(cifra.monto_comparativo or 0)

    if not has_explicit_total:
        total_row_index = len(table.rows)
        total_row = table.add_row().cells
        total_row[0].text = "Total"
        total_row[1].text = _format_note_figure_amount(nota_doc, total_actual, total_format, currency_symbol)
        total_row[2].text = ""
        total_row[3].text = _format_note_figure_amount(nota_doc, total_comparativo, total_format, currency_symbol)
        total_rows.append(total_row_index)

    _set_table_column_widths(table, [9.5, 3.2, 0.1, 3.2])
    _style_note_table(table, total_rows=total_rows, subtotal_rows=subtotal_rows, font_size=note_font_size)
    for line_index, cifra in rendered_rows:
        _apply_note_figure_format(table.rows[line_index], cifra, font_size=note_font_size)
    _force_table_font_size(table, note_font_size)

def _render_complex_note_content(document, nota_doc, labels, package, currency_symbol):
    Document, WD_ALIGN_PARAGRAPH, WD_ORIENTATION, WD_SECTION_START, WD_TABLE_ALIGNMENT, OxmlElement, qn, Cm, Pt, RGBColor = _docx_imports()
    note_font_size = nota_doc.get_print_font_size() if hasattr(nota_doc, "get_print_font_size") else BODY_SIZE
    note_alignment = nota_doc.get_print_table_alignment() if hasattr(nota_doc, "get_print_table_alignment") else "Centro"
    sections = _get_complex_note_sections(nota_doc.name)
    if not sections:
        _render_simple_note_content(document, nota_doc, labels, currency_symbol)
        return

    rendered_narrative = _get_rendered_note_narrative(nota_doc)
    if rendered_narrative:
        _add_rich_block(document, rendered_narrative, size=BODY_SIZE)
        spacer = document.add_paragraph(" ")
        _set_paragraph_runs_font(spacer, size=BODY_SIZE)

    has_visible_figures = any(
        not cint(getattr(row, "no_imprimir", 0))
        for row in (nota_doc.cifras_nota or [])
    )
    _render_note_figures(document, nota_doc, labels, note_font_size, note_alignment, currency_symbol)
    if has_visible_figures:
        spacer = document.add_paragraph("")
        _set_paragraph_runs_font(spacer, size=note_font_size)

    for section_doc in sections:
        if cint(getattr(section_doc, "mostrar_titulo", 1)):
            title_paragraph = document.add_paragraph(
                cstr(section_doc.titulo_seccion or section_doc.codigo_seccion or "Seccion"),
                style="Heading 3",
            )
            _set_paragraph_runs_font(title_paragraph, size=BODY_SIZE, bold=True)
            for run in title_paragraph.runs:
                run.bold = True
                run.italic = True
                run.underline = True

        rendered_section_narrative = _get_rendered_section_narrative(section_doc)
        if rendered_section_narrative:
            _add_rich_block(document, rendered_section_narrative, size=BODY_SIZE)

        tables = build_complex_section_tables(section_doc)
        section_table_started = False
        section_table_rendered = False
        for table_meta in tables:
            if not table_meta.get("columnas") or not table_meta.get("filas"):
                continue
            if not section_table_started:
                table_spacer = document.add_paragraph(" ")
                _set_paragraph_runs_font(table_spacer, size=note_font_size)
                section_table_started = True
            section_table_rendered = True
            compact = len(table_meta["columnas"]) > 5
            if compact:
                landscape_section = document.add_section(WD_SECTION_START.NEW_PAGE)
                _configure_section(landscape_section, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, landscape=True, document_title=REPORT_TITLE)
                _set_section_footer_page_number(landscape_section)
            data_col_indexes = [1 + (idx * 2) for idx in range(len(table_meta["columnas"]))]
            gap_col_indexes = [2 + (idx * 2) for idx in range(max(len(table_meta["columnas"]) - 1, 0))]
            header_rows = 2 if table_meta["tiene_grupos"] else 1
            table = document.add_table(rows=header_rows, cols=(len(table_meta["columnas"]) * 2))
            _set_word_table_alignment(table, note_alignment)
            _set_table_column_widths(
                table,
                _get_complex_table_widths_cm(
                    document.sections[-1],
                    len(table_meta["columnas"]),
                    gap_width_cm=0.1,
                    compact=compact,
                ),
            )
            if table_meta["tiene_grupos"]:
                top_header = table.rows[0].cells
                bottom_header = table.rows[1].cells
                top_header[0].merge(bottom_header[0]).text = "Concepto"
                cursor = 1
                for group_index, group in enumerate(table_meta["grupos_columnas"]):
                    group_width = max((cint(group.get("span", 0) or 0) * 2) - 1, 1)
                    target_cell = top_header[cursor]
                    if group_width > 1:
                        target_cell = target_cell.merge(top_header[cursor + group_width - 1])
                    target_cell.text = cstr(group["label"] or "")
                    if group_index < (len(table_meta["grupos_columnas"]) - 1):
                        top_header[cursor + group_width].text = ""
                        cursor += group_width + 1
                    else:
                        cursor += group_width
                for col_index, col in enumerate(table_meta["columnas"]):
                    cell = bottom_header[data_col_indexes[col_index]]
                    cell.text = cstr(col["etiqueta"] or col["codigo_columna"] or "-")
                    align_val = col.get("alineacion_etiqueta", "Center")
                    cell.paragraphs[0].alignment = 0 if align_val == "Left" else (2 if align_val == "Right" else 1)
            else:
                header = table.rows[0].cells
                header[0].text = "Concepto"
                for col_index, col in enumerate(table_meta["columnas"]):
                    cell = header[data_col_indexes[col_index]]
                    cell.text = cstr(col["etiqueta"] or col["codigo_columna"] or "-")
                    align_val = col.get("alineacion_etiqueta", "Center")
                    cell.paragraphs[0].alignment = 0 if align_val == "Left" else (2 if align_val == "Right" else 1)

            rendered_rows = []
            total_rows = []
            subtotal_rows = []
            for fila in table_meta["filas"]:
                row = table.add_row().cells
                row[0].text = f"{'   ' * max(cint(fila.get('nivel', 1)) - 1, 0)}{cstr(fila.get('descripcion') or fila.get('codigo_fila') or '-')}"
                for data_cell_index, cell_meta in enumerate(fila["celdas"]):
                    value_text = cstr(cell_meta.get("texto") or "")
                    if cell_meta.get("es_moneda") and value_text and value_text != "-":
                        value_text = _apply_currency_symbol(value_text, currency_symbol)
                    row[data_col_indexes[data_cell_index]].text = value_text
                row_index = len(table.rows) - 1
                rendered_rows.append((row_index, fila))
                if fila["es_total"]:
                    total_rows.append(row_index)
                elif fila["es_subtotal"]:
                    subtotal_rows.append(row_index)

            _style_note_table(
                table,
                compact=compact,
                total_rows=total_rows,
                subtotal_rows=subtotal_rows,
                header_rows=header_rows,
                font_size=min(note_font_size, COMPLEX_NOTE_TABLE_SIZE) if compact else note_font_size,
                numeric_col_indexes=tuple(data_col_indexes),
                gap_col_indexes=tuple(gap_col_indexes),
            )
            for row_index, fila in rendered_rows:
                desc_paragraph = table.rows[row_index].cells[0].paragraphs[0]
                _set_paragraph_runs_font(
                    desc_paragraph,
                    size=min(note_font_size, COMPLEX_NOTE_TABLE_SIZE) if compact else note_font_size,
                    bold=bool(
                        cint(fila.get("negrita", 0))
                        or fila["es_total"]
                        or fila["es_subtotal"]
                        or fila["es_titulo"]
                    ),
                )
                if cint(fila.get("subrayado", 0)):
                    for run in desc_paragraph.runs:
                        run.underline = True
                for data_cell_index, cell_meta in enumerate(fila["celdas"]):
                    paragraph = table.rows[row_index].cells[data_col_indexes[data_cell_index]].paragraphs[0]
                    if cell_meta["alineacion"] == "Left":
                        paragraph.alignment = 0
                    elif cell_meta["alineacion"] == "Center":
                        paragraph.alignment = 1
                    else:
                        paragraph.alignment = 2
            _force_table_font_size(
                table,
                min(note_font_size, COMPLEX_NOTE_TABLE_SIZE) if compact else note_font_size,
            )

            if compact:
                portrait_section = document.add_section(WD_SECTION_START.NEW_PAGE)
                _configure_section(portrait_section, package, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, landscape=False, document_title=REPORT_TITLE)
                _set_section_footer_page_number(portrait_section)

        if section_table_rendered:
            table_end_spacer = document.add_paragraph(" ")
            _set_paragraph_runs_font(table_end_spacer, size=note_font_size)

        if section_doc.observaciones:
            paragraph = document.add_paragraph(cstr(section_doc.observaciones))
            _set_paragraph_runs_font(paragraph, size=BODY_SIZE)


def _get_complex_note_sections(note_name):
    section_rows = frappe.get_all(
        "Seccion Nota EEFF",
        filters={"nota_eeff": note_name},
        fields=["name"],
        order_by="orden asc, creation asc",
        limit_page_length=500,
    )
    return [frappe.get_doc("Seccion Nota EEFF", row.name) for row in section_rows]


def _build_complex_note_tables(cells):
    tables = {}
    for cell in cells or []:
        table_code = cstr(getattr(cell, "codigo_tabla", "")).strip().upper() or "TABLA_01"
        row_code = cstr(getattr(cell, "codigo_fila", "")).strip().upper() or "FILA"
        col_code = cstr(getattr(cell, "codigo_columna", "")).strip().upper() or "COLUMNA"

        table_meta = tables.setdefault(
            table_code,
            {
                "codigo_tabla": table_code,
                "columnas": [],
                "filas": [],
                "valores": {},
            },
        )
        if col_code not in table_meta["columnas"]:
            table_meta["columnas"].append(col_code)
        if row_code not in table_meta["filas"]:
            table_meta["filas"].append(row_code)

        value_text = cstr(getattr(cell, "valor_texto", "")).strip()
        value_number = getattr(cell, "valor_numero", None)
        if value_text:
            value = value_text
        elif value_number not in (None, ""):
            value = _fmt_number(value_number)
        else:
            value = "-"
        table_meta["valores"][(row_code, col_code)] = value

    output = []
    for key in sorted(tables.keys()):
        table_meta = tables[key]
        table_meta["columnas"] = sorted(table_meta["columnas"])
        table_meta["filas"] = sorted(table_meta["filas"])
        output.append(table_meta)
    return output


def _add_datos_estadisticos_section(document, package):
    document.add_paragraph("Datos Estadisticos", style="Heading 1")
    labels = package.get_column_labels()
    comparative_map = package.get_datos_estadisticos_comparativos_map() if hasattr(package, "get_datos_estadisticos_comparativos_map") else {}
    datos = sorted(
        [row for row in (package.datos_estadisticos or []) if not cint(getattr(row, "no_imprimir", 0))],
        key=lambda row: (cint(getattr(row, "orden", 0)), cint(getattr(row, "idx", 0))),
    )
    if not datos:
        paragraph = document.add_paragraph("No hay datos estadisticos registrados para este paquete.")
        _set_paragraph_runs_font(paragraph)
        return

    table = document.add_table(rows=1, cols=4)
    headers = ["Dato", "Unidad", labels["actual"], labels["comparativo"]]
    for header_index, title in enumerate(headers):
        table.rows[0].cells[header_index].text = title

    for row in datos:
        code = cstr(getattr(row, "codigo_dato", "") or "").strip().upper()
        data_row = table.add_row().cells
        data_row[0].text = cstr(row.descripcion or row.codigo_dato or "-")
        data_row[1].text = cstr(row.unidad_medida or "-")
        data_row[2].text = _fmt_number(row.valor_actual)
        data_row[3].text = _fmt_number(comparative_map.get(code, 0))

    _set_table_column_widths(table, [7.0, 2.5, 3.0, 3.0])
    _style_note_table(table)


def _build_tabular_sections(celdas):
    sections = {}
    for celda in celdas or []:
        seccion_id = cstr(getattr(celda, "seccion_id", "")).strip().upper() or "GENERAL"
        codigo_fila = cstr(getattr(celda, "codigo_fila", "")).strip().upper() or "FILA"
        codigo_columna = cstr(getattr(celda, "codigo_columna", "")).strip().upper() or "COLUMNA"

        sec = sections.setdefault(
            seccion_id,
            {
                "seccion_id": seccion_id,
                "columnas": [],
                "filas": [],
                "valores": {},
            },
        )
        if codigo_columna not in sec["columnas"]:
            sec["columnas"].append(codigo_columna)
        if codigo_fila not in sec["filas"]:
            sec["filas"].append(codigo_fila)

        valor_texto = cstr(getattr(celda, "valor_texto", "")).strip()
        valor_numero = getattr(celda, "valor_numero", None)
        if valor_texto:
            valor = valor_texto
        elif valor_numero not in (None, ""):
            valor = _fmt_number(valor_numero)
        else:
            valor = "-"

        sec["valores"][(codigo_fila, codigo_columna)] = valor

    output = []
    for key in sorted(sections.keys()):
        section = sections[key]
        section["columnas"] = sorted(section["columnas"])
        section["filas"] = sorted(section["filas"])
        output.append(section)
    return output


def _apply_note_figure_format(row, cifra, font_size=BODY_SIZE):
    desc_paragraph = row.cells[0].paragraphs[0]
    _set_paragraph_runs_font(
        desc_paragraph,
        size=font_size,
        bold=bool(
            cint(getattr(cifra, "negrita", 0))
            or cint(getattr(cifra, "es_titulo", 0))
            or cint(getattr(cifra, "es_total", 0))
            or cint(getattr(cifra, "es_subtotal", 0))
        ),
    )
    if cint(getattr(cifra, "subrayado", 0)):
        for run in desc_paragraph.runs:
            run.underline = True
    if cint(getattr(cifra, "es_linea_blanco", 0)):
        _set_row_min_height(row, 340)


def _fill_label_value_cell(cell, label, value):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.add_run(f"{label}\n").bold = True
    paragraph.add_run(cstr(value or "-"))
    _set_paragraph_runs_font(paragraph)


def _set_paragraph_runs_font(paragraph, font_name=FONT_NAME, size=BODY_SIZE, bold=False):
    _Document, _Align, _Orientation, _SectionStart, _TableAlign, OxmlElement, qn, _Cm, Pt, RGBColor = _docx_imports()
    if not paragraph.runs:
        paragraph.add_run("")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        if bold:
            run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        rfonts.set(qn("w:eastAsia"), font_name)


def _set_table_column_widths(table, widths_cm):
    _Document, _Align, _Orientation, _SectionStart, _TableAlign, OxmlElement, qn, Cm, _Pt, _RGBColor = _docx_imports()
    table.autofit = False

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_twips = sum(int(Cm(width_cm).emu / 635) for width_cm in widths_cm)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_twips))

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width_cm in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(Cm(width_cm).emu / 635)))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, width_cm in enumerate(widths_cm):
            if index >= len(row.cells):
                break
            cell = row.cells[index]
            cell.width = Cm(width_cm)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(Cm(width_cm).emu / 635)))


def _set_word_table_alignment(table, alignment):
    _Document, _Align, _Orientation, _SectionStart, WD_TABLE_ALIGNMENT, _OxmlElement, _qn, _Cm, _Pt, _RGBColor = _docx_imports()
    if alignment == "Izquierda":
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
    elif alignment == "Derecha":
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    else:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER


def _force_table_font_size(table, font_size):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _set_paragraph_runs_font(paragraph, size=font_size)


def _get_estado_table_widths_cm(section, estado_doc):
    # Word usa anchos fijos en cm para evitar que el motor DOCX redistribuya
    # de forma inconsistente la columna comparativa.
    return [10.0, 1.25, 2.9, 0.2, 2.9]


def _get_complex_table_widths_cm(section, data_columns, gap_width_cm=0.1, compact=False):
    data_columns = max(cint(data_columns or 0), 1)
    gap_count = max(data_columns - 1, 0)
    usable_width_cm = max(_section_usable_width_cm(section), 0)
    gap_total_cm = gap_count * flt(gap_width_cm or 0.1)

    # Intenta fijar 3cm por columna de datos; si no cabe, reduce de forma controlada.
    base_concept_cm = 7.0 if compact else 9.5
    target_data_cm = 3.0
    min_data_cm = 1.0 if compact else 1.2
    min_concept_cm = 4.0

    remaining_after_target = usable_width_cm - gap_total_cm - (data_columns * target_data_cm)
    if remaining_after_target >= min_concept_cm:
        concept_cm = min(base_concept_cm, remaining_after_target)
        data_cm = target_data_cm
    else:
        concept_cm = min(base_concept_cm, max(remaining_after_target, min_concept_cm))
        available_for_data_cm = max(usable_width_cm - gap_total_cm - concept_cm, data_columns * min_data_cm)
        data_cm = max(available_for_data_cm / data_columns, min_data_cm)

    widths = [concept_cm]
    for index in range(data_columns):
        widths.append(data_cm)
        if index < data_columns - 1:
            widths.append(flt(gap_width_cm or 0.1))
    return widths


def _section_usable_width_cm(section):
    page_width_emu = float(section.page_width)
    left_margin_emu = float(section.left_margin)
    right_margin_emu = float(section.right_margin)
    usable_emu = max(page_width_emu - left_margin_emu - right_margin_emu, 0)
    return usable_emu / 360000.0


def _parse_word_table_width_percent(value):
    raw = cstr(value or "").strip()
    match = re.fullmatch(r"(\d+(?:\.\d+)?)%", raw)
    if not match:
        return 1.0
    percent = flt(match.group(1)) / 100.0
    return max(0.4, min(percent, 1.0))


def _style_table_no_borders(table):
    _Document, _Align, _Orientation, _SectionStart, _TableAlign, OxmlElement, qn, _Cm, _Pt, _RGBColor = _docx_imports()
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "nil")


def _style_cell_border(cell, top=None, bottom=None):
    _Document, _Align, _Orientation, _SectionStart, _TableAlign, OxmlElement, qn, _Cm, _Pt, _RGBColor = _docx_imports()
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, spec in (("top", top), ("bottom", bottom)):
        if not spec:
            continue
        edge_el = tcBorders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tcBorders.append(edge_el)
        edge_el.set(qn("w:val"), spec.get("val", "single"))
        edge_el.set(qn("w:sz"), str(spec.get("sz", 8)))
        edge_el.set(qn("w:color"), spec.get("color", "000000"))


def _set_cell_margins(cell, left=None, right=None):
    _Document, _Align, _Orientation, _SectionStart, _TableAlign, OxmlElement, qn, _Cm, _Pt, _RGBColor = _docx_imports()
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for edge, value in (("left", left), ("right", right)):
        if value is None:
            continue
        edge_el = tcMar.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tcMar.append(edge_el)
        edge_el.set(qn("w:w"), str(value))
        edge_el.set(qn("w:type"), "dxa")


def _set_row_min_height(row, height_twips):
    _Document, _Align, _Orientation, _SectionStart, _TableAlign, OxmlElement, qn, _Cm, _Pt, _RGBColor = _docx_imports()
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.first_child_found_in("w:trHeight")
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height_twips))
    tr_height.set(qn("w:hRule"), "atLeast")


def _style_financial_table(
    table,
    compact=False,
    total_rows=None,
    subtotal_rows=None,
    header_rows=1,
    font_size=None,
    note_col_index=None,
    numeric_col_indexes=(1, 3),
    gap_col_index=2,
    gap_col_indexes=None,
    align_numeric_headers=False,
):
    total_rows = total_rows or []
    subtotal_rows = subtotal_rows or []
    effective_font_size = font_size or (COMPACT_SIZE if compact else BODY_SIZE)
    numeric_indexes = tuple(sorted({cint(idx) for idx in (numeric_col_indexes or [])}))
    if gap_col_indexes is None:
        gap_indexes = {cint(gap_col_index)} if gap_col_index is not None else set()
    else:
        gap_indexes = {cint(idx) for idx in (gap_col_indexes or [])}
    first_numeric_index = min(numeric_indexes) if numeric_indexes else None
    _style_table_no_borders(table)
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            if note_col_index is not None and cell_index == note_col_index:
                _set_cell_margins(cell, left=80, right=80)
            elif first_numeric_index is not None and cell_index == first_numeric_index:
                _set_cell_margins(cell, left=140, right=100)
            elif cell_index in gap_indexes:
                _set_cell_margins(cell, left=0, right=0)
            elif cell_index in numeric_indexes:
                _set_cell_margins(cell, left=140, right=100)
            for paragraph in cell.paragraphs:
                _set_paragraph_runs_font(
                    paragraph,
                    size=effective_font_size,
                    bold=(row_index < header_rows or row_index in total_rows or row_index in subtotal_rows),
                )
                if cell_index in numeric_indexes and (
                    row_index >= header_rows
                    or (align_numeric_headers and row_index == max(header_rows - 1, 0))
                ):
                    paragraph.alignment = 2
                elif note_col_index is not None and cell_index == note_col_index:
                    paragraph.alignment = 1
                elif cell_index in gap_indexes:
                    paragraph.alignment = 0
        if row_index in subtotal_rows:
            for cell_index, cell in enumerate(row.cells):
                if cell_index in numeric_indexes:
                    _style_cell_border(cell, top={"val": "single", "sz": 8}, bottom={"val": "single", "sz": 8})
        if row_index in total_rows:
            for cell_index, cell in enumerate(row.cells):
                if cell_index in numeric_indexes:
                    _style_cell_border(cell, top={"val": "single", "sz": 8}, bottom={"val": "double", "sz": 12})


def _style_note_table(
    table,
    compact=False,
    total_rows=None,
    subtotal_rows=None,
    header_rows=1,
    font_size=None,
    numeric_col_indexes=(1, 3),
    gap_col_indexes=(2,),
):
    _style_financial_table(
        table,
        compact=compact,
        total_rows=total_rows,
        subtotal_rows=subtotal_rows,
        header_rows=header_rows,
        font_size=font_size,
        numeric_col_indexes=numeric_col_indexes,
        gap_col_indexes=gap_col_indexes,
        align_numeric_headers=True,
    )


def _style_meta_table(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _set_paragraph_runs_font(paragraph)


def _add_rich_block(document, raw_value, size=BODY_SIZE):
    html = cstr(raw_value or "")
    if not html.strip():
        return

    segments = _split_html_with_tables(html)
    for segment in segments:
        if segment.get("type") == "table":
            rendered = _add_html_table_block(document, cstr(segment.get("content") or ""), size=size)
            if rendered:
                spacer = document.add_paragraph()
                _set_paragraph_runs_font(spacer, size=size)
            continue
        _add_rich_text_block(document, cstr(segment.get("content") or ""), size=size)


def _add_rich_text_block(document, raw_html, size=BODY_SIZE):
    blocks = _html_to_word_blocks(raw_html)
    if not blocks:
        text = _html_to_text(raw_html)
        if not text:
            return
        for block in [item.strip() for item in re.split(r"\n\s*\n", text) if item.strip()]:
            paragraph = document.add_paragraph(block)
            _set_paragraph_runs_font(paragraph, size=size)
        return

    for block in blocks:
        block_text = "".join(cstr(run.get("text") or "") for run in block)
        paragraph = document.add_paragraph()
        if not block_text.replace("\n", "").strip():
            _set_paragraph_runs_font(paragraph, size=size)
            continue
        _append_runs_to_paragraph(paragraph, block)
        _set_paragraph_runs_font(paragraph, size=size)


def _split_html_with_tables(html):
    text = cstr(html or "")
    if not re.search(r"<table\b", text, flags=re.IGNORECASE):
        return [{"type": "html", "content": text}]

    parts = re.split(r"(<table[\s\S]*?</table>)", text, flags=re.IGNORECASE)
    segments = []
    for part in parts:
        if not cstr(part or "").strip():
            continue
        if re.match(r"^\s*<table\b", part, flags=re.IGNORECASE):
            segments.append({"type": "table", "content": part})
        else:
            segments.append({"type": "html", "content": part})
    return segments or [{"type": "html", "content": text}]


def _add_html_table_block(document, raw_table_html, size=BODY_SIZE):
    rows = _html_table_to_rows(raw_table_html)
    if not rows:
        return False

    col_count = max((len(row) for row in rows), default=0)
    if col_count <= 0:
        return False

    table = document.add_table(rows=len(rows), cols=col_count)
    table.autofit = True
    _style_table_no_borders(table)

    for row_idx, row in enumerate(rows):
        for col_idx in range(col_count):
            cell_meta = row[col_idx] if col_idx < len(row) else {"runs": [], "is_header": False}
            paragraph = table.rows[row_idx].cells[col_idx].paragraphs[0]
            paragraph.text = ""
            _append_runs_to_paragraph(paragraph, cell_meta.get("runs") or [])
            _set_paragraph_runs_font(paragraph, size=size, bold=bool(cell_meta.get("is_header")))
    return True


def _html_table_to_rows(raw_table_html):
    parser = _HtmlTableParser()
    parser.feed(cstr(raw_table_html or ""))
    parser.close()
    return parser.rows


def _append_runs_to_paragraph(paragraph, runs_meta):
    for run_meta in runs_meta or []:
        parts = cstr(run_meta.get("text") or "").split("\n")
        for part_index, part in enumerate(parts):
            run = paragraph.add_run(part)
            if run_meta.get("bold"):
                run.bold = True
            if run_meta.get("italic"):
                run.italic = True
            if run_meta.get("underline"):
                run.underline = True
            _apply_run_background(run, run_meta.get("bg_color"))
            if part_index < len(parts) - 1:
                run.add_break()


def _apply_run_background(run, color_hex):
    fill = cstr(color_hex or "").strip().upper()
    if not re.fullmatch(r"[0-9A-F]{6}", fill):
        return
    _Document, _Align, _Orientation, _SectionStart, _TableAlign, OxmlElement, qn, _Cm, _Pt, _RGBColor = _docx_imports()
    rpr = run._element.get_or_add_rPr()
    shd = rpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        rpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _derive_text_style(base_style, tag, attrs_map, bold_tags, italic_tags, underline_tags):
    style = dict(base_style or {})
    clean_tag = cstr(tag or "").lower()
    style_text = cstr(attrs_map.get("style") or "")
    style_text_compact = style_text.lower().replace(" ", "")

    if clean_tag in set(bold_tags or []):
        style["bold"] = True
    if clean_tag in set(italic_tags or []):
        style["italic"] = True
    if clean_tag in set(underline_tags or []):
        style["underline"] = True

    if "font-weight:bold" in style_text_compact:
        style["bold"] = True
    for weight in ("font-weight:600", "font-weight:700", "font-weight:800", "font-weight:900"):
        if weight in style_text_compact:
            style["bold"] = True
            break
    if "font-style:italic" in style_text_compact:
        style["italic"] = True
    if "text-decoration:underline" in style_text_compact or "text-decoration-line:underline" in style_text_compact:
        style["underline"] = True

    bg_color = _extract_bg_color_hex(attrs_map)
    if clean_tag == "mark" and not bg_color:
        bg_color = "FFFF00"
    if bg_color:
        style["bg_color"] = bg_color

    return style


def _extract_bg_color_hex(attrs_map):
    style_text = cstr((attrs_map or {}).get("style") or "")
    bg_color_raw = _extract_css_property(style_text, "background-color")
    if not bg_color_raw:
        bg_shorthand = _extract_css_property(style_text, "background")
        if bg_shorthand and "url(" not in bg_shorthand.lower():
            tokens = [token.strip() for token in bg_shorthand.split() if token.strip()]
            if tokens:
                bg_color_raw = tokens[-1]

    color_hex = _css_color_to_hex(bg_color_raw)
    if color_hex:
        return color_hex
    return _css_color_to_hex(cstr((attrs_map or {}).get("bgcolor") or ""))


def _extract_css_property(style_text, property_name):
    if not style_text or not property_name:
        return ""
    pattern = rf"(?:^|;)\s*{re.escape(property_name)}\s*:\s*([^;]+)"
    match = re.search(pattern, cstr(style_text), flags=re.IGNORECASE)
    if not match:
        return ""
    return cstr(match.group(1) or "").strip()


def _css_color_to_hex(raw_color):
    value = cstr(raw_color or "").strip().lower()
    if not value or value in ("transparent", "none", "inherit", "initial"):
        return ""

    named = {
        "black": "000000",
        "white": "FFFFFF",
        "gray": "808080",
        "grey": "808080",
        "lightgray": "D3D3D3",
        "lightgrey": "D3D3D3",
        "silver": "C0C0C0",
        "yellow": "FFFF00",
        "red": "FF0000",
        "green": "008000",
        "blue": "0000FF",
        "cyan": "00FFFF",
        "magenta": "FF00FF",
        "orange": "FFA500",
        "pink": "FFC0CB",
        "purple": "800080",
        "brown": "A52A2A",
    }
    if value in named:
        return named[value]

    hex_match = re.fullmatch(r"#([0-9a-f]{3}|[0-9a-f]{6})", value, flags=re.IGNORECASE)
    if hex_match:
        token = hex_match.group(1)
        if len(token) == 3:
            token = "".join(ch * 2 for ch in token)
        return token.upper()

    rgb_match = re.fullmatch(r"rgba?\(([^)]+)\)", value)
    if rgb_match:
        pieces = [chunk.strip() for chunk in rgb_match.group(1).split(",")]
        if len(pieces) >= 3:
            rgb_values = []
            for piece in pieces[:3]:
                if piece.endswith("%"):
                    percent = flt(piece[:-1])
                    channel = int(max(0, min(255, round((percent / 100.0) * 255))))
                else:
                    channel = int(max(0, min(255, round(flt(piece)))))
                rgb_values.append(channel)
            return "".join(f"{channel:02X}" for channel in rgb_values)
    return ""


class _HtmlTableParser(HTMLParser):
    VOID_TAGS = {"br", "hr", "img"}
    BOLD_TAGS = {"b", "strong", "th", "h1", "h2", "h3", "h4", "h5", "h6"}
    ITALIC_TAGS = {"i", "em"}
    UNDERLINE_TAGS = {"u"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.rows = []
        self._current_row = None
        self._current_cell = None
        self._in_cell = False
        self._style_stack = [{"bold": False, "italic": False, "underline": False, "bg_color": ""}]

    def handle_starttag(self, tag, attrs):
        tag = cstr(tag or "").lower()
        attrs_map = {cstr(key or "").lower(): cstr(value or "") for key, value in (attrs or [])}

        if tag == "tr":
            self._current_row = []
            return

        if tag in ("td", "th") and self._current_row is not None:
            colspan = cint(attrs_map.get("colspan") or 1)
            colspan = max(colspan, 1)
            self._current_cell = {"runs": [], "is_header": tag == "th", "colspan": colspan}
            self._in_cell = True

        if tag == "br" and self._in_cell and self._current_cell is not None:
            self._append_cell_text("\n", self._style_stack[-1])
            return

        if tag == "li" and self._in_cell and self._current_cell is not None:
            self._append_cell_text("- ", self._style_stack[-1])

        if tag in self.VOID_TAGS:
            return

        next_style = _derive_text_style(
            self._style_stack[-1],
            tag,
            attrs_map,
            self.BOLD_TAGS,
            self.ITALIC_TAGS,
            self.UNDERLINE_TAGS,
        )
        if tag == "th":
            next_style["bold"] = True
        self._style_stack.append(next_style)

    def handle_endtag(self, tag):
        tag = cstr(tag or "").lower()

        if tag in ("td", "th") and self._in_cell and self._current_row is not None and self._current_cell is not None:
            cell_meta = {
                "runs": self._normalize_cell_runs(self._current_cell.get("runs") or []),
                "is_header": bool(self._current_cell.get("is_header")),
            }
            self._current_row.append(cell_meta)
            for _idx in range(max(cint(self._current_cell.get("colspan") or 1) - 1, 0)):
                self._current_row.append({"runs": [], "is_header": bool(self._current_cell.get("is_header"))})
            self._current_cell = None
            self._in_cell = False
            if len(self._style_stack) > 1:
                self._style_stack.pop()
            return

        if tag == "tr" and self._current_row is not None:
            if self._current_row:
                self.rows.append(self._current_row)
            self._current_row = None
            return

        if tag in self.VOID_TAGS:
            return

        if len(self._style_stack) > 1:
            self._style_stack.pop()

    def handle_data(self, data):
        if not self._in_cell or self._current_cell is None:
            return
        self._append_cell_text(cstr(data or "").replace("\xa0", " "), self._style_stack[-1])

    def _append_cell_text(self, text, style):
        clean_text = cstr(text or "")
        if not clean_text or self._current_cell is None:
            return
        runs = self._current_cell["runs"]
        if runs and self._same_style(runs[-1], style):
            runs[-1]["text"] += clean_text
            return
        runs.append(
            {
                "text": clean_text,
                "bold": bool(style.get("bold")),
                "italic": bool(style.get("italic")),
                "underline": bool(style.get("underline")),
                "bg_color": cstr(style.get("bg_color") or "").upper(),
            }
        )

    @staticmethod
    def _same_style(run_meta, style):
        return (
            bool(run_meta.get("bold")) == bool(style.get("bold"))
            and bool(run_meta.get("italic")) == bool(style.get("italic"))
            and bool(run_meta.get("underline")) == bool(style.get("underline"))
            and cstr(run_meta.get("bg_color") or "").upper() == cstr(style.get("bg_color") or "").upper()
        )

    @staticmethod
    def _normalize_cell_runs(runs):
        normalized = []
        for run in runs:
            text = cstr(run.get("text") or "")
            if not text:
                continue
            text = re.sub(r"[ \t]{2,}", " ", text)
            normalized.append(
                {
                    "text": text,
                    "bold": bool(run.get("bold")),
                    "italic": bool(run.get("italic")),
                    "underline": bool(run.get("underline")),
                    "bg_color": cstr(run.get("bg_color") or "").upper(),
                }
            )
        return normalized


def _html_to_word_blocks(value):
    parser = _WordRichTextParser()
    parser.feed(cstr(value or ""))
    parser.close()
    return parser.get_blocks()


class _WordRichTextParser(HTMLParser):
    BLOCK_TAGS = {"p", "li", "h1", "h2", "h3", "h4", "h5", "h6"}
    VOID_TAGS = {"br", "hr", "img"}
    BOLD_TAGS = {"b", "strong", "h1", "h2", "h3", "h4", "h5", "h6"}
    ITALIC_TAGS = {"i", "em"}
    UNDERLINE_TAGS = {"u"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self._blocks = []
        self._current_block = None
        self._style_stack = [{"bold": False, "italic": False, "underline": False, "bg_color": ""}]

    def handle_starttag(self, tag, attrs):
        tag = cstr(tag or "").lower()
        attrs_map = {cstr(key or "").lower(): cstr(value or "") for key, value in (attrs or [])}

        if tag in self.BLOCK_TAGS:
            self._start_block()
            if tag == "li":
                self._append_text("- ", self._style_stack[-1])

        if tag == "br":
            self._ensure_block()
            self._append_text("\n", self._style_stack[-1])
            return

        if tag in self.VOID_TAGS:
            return

        self._style_stack.append(self._derive_style(tag, attrs_map))

    def handle_endtag(self, tag):
        tag = cstr(tag or "").lower()
        if tag in self.BLOCK_TAGS:
            self._finalize_block()

        if tag in self.VOID_TAGS:
            return

        if len(self._style_stack) > 1:
            self._style_stack.pop()

    def handle_data(self, data):
        text = cstr(data or "").replace("\xa0", " ")
        if not text:
            return
        self._ensure_block()
        self._append_text(text, self._style_stack[-1])

    def get_blocks(self):
        self._finalize_block()
        return [block for block in self._blocks if self._has_visible_content(block)]

    def _derive_style(self, tag, attrs_map):
        return _derive_text_style(
            self._style_stack[-1],
            tag,
            attrs_map,
            self.BOLD_TAGS,
            self.ITALIC_TAGS,
            self.UNDERLINE_TAGS,
        )

    def _ensure_block(self):
        if self._current_block is None:
            self._current_block = []

    def _start_block(self):
        self._finalize_block()
        self._current_block = []

    def _finalize_block(self):
        if self._current_block is None:
            return
        self._blocks.append(self._current_block)
        self._current_block = None

    def _append_text(self, text, style):
        if not text:
            return
        if self._current_block and self._same_style(self._current_block[-1], style):
            self._current_block[-1]["text"] += text
            return
        self._current_block.append(
            {
                "text": text,
                "bold": bool(style.get("bold")),
                "italic": bool(style.get("italic")),
                "underline": bool(style.get("underline")),
                "bg_color": cstr(style.get("bg_color") or "").upper(),
            }
        )

    @staticmethod
    def _same_style(run_meta, style):
        return (
            bool(run_meta.get("bold")) == bool(style.get("bold"))
            and bool(run_meta.get("italic")) == bool(style.get("italic"))
            and bool(run_meta.get("underline")) == bool(style.get("underline"))
            and cstr(run_meta.get("bg_color") or "").upper() == cstr(style.get("bg_color") or "").upper()
        )

    @staticmethod
    def _has_visible_content(block):
        if not block:
            return False
        raw_text = "".join(cstr(run.get("text") or "") for run in block)
        if raw_text.replace("\n", "").strip():
            return True
        return "\n" in raw_text


def _html_to_text(value):
    text = cstr(value or "")
    if not text.strip():
        return ""
    replacements = (
        (r"<br\s*/?>", "\n"),
        (r"</p>", "\n\n"),
        (r"</div>", "\n"),
        (r"</li>", "\n"),
        (r"<li[^>]*>", "- "),
        (r"</h[1-6]>", "\n"),
    )
    for pattern, repl in replacements:
        text = re.sub(pattern, repl, text, flags=re.IGNORECASE)
    text = re.sub(r"<[^>]+>", "", text)
    text = unescape(text)
    text = text.replace("\xa0", " ")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _package_has_signatures(package):
    return any(
        cstr(getattr(package, field, "") or "").strip()
        for field in ("elaborado_por", "revisado_por", "autorizado_por")
    )


def _add_package_signatures_block(document, package):
    table = document.add_table(rows=2, cols=3)
    table.autofit = False
    _set_table_column_widths(table, [5.5, 5.5, 5.5])
    _style_table_no_borders(table)

    roles = [
        (cstr(getattr(package, "elaborado_por", "") or "").strip(), cstr(getattr(package, "cargo_elaborado_por", "") or "").strip()),
        (cstr(getattr(package, "revisado_por", "") or "").strip(), cstr(getattr(package, "cargo_revisado_por", "") or "").strip()),
        (cstr(getattr(package, "autorizado_por", "") or "").strip(), cstr(getattr(package, "cargo_autorizado_por", "") or "").strip()),
    ]

    for idx, (name, title) in enumerate(roles):
        top_cell = table.rows[0].cells[idx]
        bottom_cell = table.rows[1].cells[idx]
        top_cell.text = "\n\n\n"
        _set_paragraph_runs_font(top_cell.paragraphs[0])
        if not name:
            bottom_cell.text = ""
            _set_paragraph_runs_font(bottom_cell.paragraphs[0])
            continue

        bottom_cell.text = ""
        line_paragraph = bottom_cell.paragraphs[0]
        line_paragraph.alignment = 1
        line_paragraph.add_run("")
        _set_paragraph_runs_font(line_paragraph, size=10)
        _style_cell_border(bottom_cell, top={"val": "single", "sz": 8})

        name_paragraph = bottom_cell.add_paragraph(name)
        name_paragraph.alignment = 1
        _set_paragraph_runs_font(name_paragraph, size=11)

        if title:
            title_paragraph = bottom_cell.add_paragraph(title)
            title_paragraph.alignment = 1
            _set_paragraph_runs_font(title_paragraph, size=11)


def _fmt_number(value):
    return format_accounting_number(value, "Numero", trim_plain=True, none_as="-")


def _get_package_currency_symbol(package):
    if hasattr(package, "get_currency_context"):
        ctx = package.get_currency_context() or {}
        symbol = cstr(ctx.get("symbol") or "").strip()
        if symbol:
            return symbol
    return "$"


def _is_estado_currency_line(linea):
    return cstr(getattr(linea, "formato_presentacion", "") or "").strip().title() == "Moneda"


def _is_note_currency_figure(nota_doc, cifra):
    if hasattr(nota_doc, "get_figure_number_format"):
        return cstr(nota_doc.get_figure_number_format(cifra) or "").strip().title() == "Moneda"
    return False


def _apply_currency_symbol(value_text, symbol):
    text = cstr(value_text or "").strip()
    if not text:
        return text
    return f"{symbol} {text}"


def _format_note_figure_value(nota_doc, cifra, fieldname, currency_symbol):
    if cint(getattr(cifra, "es_linea_blanco", 0)) or cint(getattr(cifra, "es_titulo", 0)):
        return ""
    output = ""
    if hasattr(nota_doc, "format_figure_value"):
        output = cstr(nota_doc.format_figure_value(cifra, fieldname) or "-")
    else:
        output = _fmt_number(getattr(cifra, fieldname, None))
    if _is_note_currency_figure(nota_doc, cifra):
        return _apply_currency_symbol(output, currency_symbol)
    return output


def _format_note_figure_amount(nota_doc, value, format_type, currency_symbol):
    output = ""
    if hasattr(nota_doc, "format_figure_amount"):
        output = cstr(nota_doc.format_figure_amount(value, format_type) or "-")
    else:
        output = _fmt_number(value)
    if cstr(format_type or "").strip().title() == "Moneda":
        return _apply_currency_symbol(output, currency_symbol)
    return output


def _is_text_note_figure(nota_doc, cifra):
    if cint(getattr(cifra, "es_linea_blanco", 0)) or cint(getattr(cifra, "es_titulo", 0)):
        return False
    if hasattr(nota_doc, "is_text_figure"):
        return bool(nota_doc.is_text_figure(cifra))
    return False


def _get_note_figures_total_format(nota_doc, figures):
    if hasattr(nota_doc, "get_figures_total_format"):
        return cstr(nota_doc.get_figures_total_format(figures) or "Moneda")
    return "Moneda"


def _get_rendered_note_observaciones(nota_doc):
    if hasattr(nota_doc, "render_observaciones"):
        output = cstr(nota_doc.render_observaciones() or "").strip()
        if output:
            return output
    return cstr(getattr(nota_doc, "observaciones", "") or "").strip()


def _get_rendered_note_narrative(nota_doc):
    if hasattr(nota_doc, "render_contenido_narrativo"):
        output = cstr(nota_doc.render_contenido_narrativo() or "").strip()
        if output:
            return output
    return cstr(getattr(nota_doc, "contenido_narrativo", "") or "").strip()


def _get_rendered_section_narrative(section_doc):
    if hasattr(section_doc, "render_contenido_narrativo"):
        output = cstr(section_doc.render_contenido_narrativo() or "").strip()
        if output:
            return output
    return cstr(getattr(section_doc, "contenido_narrativo", "") or "").strip()


def _build_word_export_filename(package_name):
    safe_name = _sanitize_filename(package_name)
    base_name = f"{REPORT_TITLE} - {safe_name}" if safe_name else REPORT_TITLE
    max_base_length = 135 - len(".docx")
    base_name = base_name[:max_base_length].rstrip(" -_")
    return f"{base_name}.docx"


def _sanitize_filename(value):
    filename = re.sub(r'[\\/:*?"<>|]+', "-", cstr(value or REPORT_TITLE)).strip()
    return filename or REPORT_TITLE


